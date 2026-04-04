"""
Create publication-ready Word document (.docx) from research paper Markdown.
Applies professional formatting for SSRN and ArXiv submission.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
from datetime import datetime

# Color scheme
HEADING_COLOR = RGBColor(25, 55, 109)  # Dark blue
ACCENT_COLOR = RGBColor(68, 114, 196)  # Steel blue

def set_cell_background(cell, fill):
    """Set background color for table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_page_break(doc):
    """Add page break"""
    doc.add_page_break()

def set_document_margins(doc, top=1.0, bottom=1.0, left=1.0, right=1.0):
    """Set document margins (in inches)"""
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)

def add_header_footer(doc, header_text, footer_text):
    """Add header and footer to all sections"""
    for section in doc.sections:
        # Header
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.text = header_text
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_para_format = header_para.paragraph_format
        header_para_format.font_size = Pt(10)
        
        # Footer with page numbers
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add page number field
        run = footer_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

def create_publication_word_doc(markdown_file, output_file):
    """Create professional Word document from Markdown"""
    
    # Read markdown content
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create document
    doc = Document()
    
    # Set margins
    set_document_margins(doc, top=1.0, bottom=1.0, left=1.0, right=1.0)
    
    # Set default font for entire document
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # Add header/footer
    add_header_footer(doc, "GlassBox: Runtime Decision Governance Framework", "")
    
    # ===== TITLE PAGE =====
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("GlassBox: A Runtime Decision Governance Framework\nfor Autonomous Enterprise AI Systems")
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = HEADING_COLOR
    
    # Spacing after title
    for _ in range(2):
        doc.add_paragraph()
    
    # Author info
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_para.add_run("Independent Researcher")
    author_run.font.size = Pt(12)
    
    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"April 2026")
    date_run.font.size = Pt(11)
    
    # Spacing
    for _ in range(2):
        doc.add_paragraph()
    
    # Abstract
    abstract_heading = doc.add_paragraph()
    abstract_heading_run = abstract_heading.add_run("Abstract")
    abstract_heading_run.font.bold = True
    abstract_heading_run.font.size = Pt(12)
    
    abstract_text = """As enterprise AI systems shift from model-centric to agent-centric architectures, organizations face a critical challenge: autonomous agents make high-stakes decisions without built-in governance mechanisms. We introduce GlassBox, a runtime decision governance framework that addresses this gap. GlassBox provides a transparent, auditable layer between AI agents and enterprise systems, offering real-time policy enforcement, anomaly detection, and risk quantification. The framework combines nine governance stages—from contract validation to audit logging—with distributed velocity breaking, multi-agent risk aggregation, and formal policy composition. We demonstrate formal correctness through mathematical modeling, validate performance through comprehensive benchmarks (achieving 95.6% fraud detection accuracy with <1% latency overhead), and show enterprise applicability across compliance domains (NIST AI RMF, GDPR, HIPAA, SOC 2, ISO 27001). GlassBox handles 333+ decisions per second in single-instance deployment and 100,000+ decisions per second in distributed mode. Our evaluation covers throughput scaling, anomaly detection precision, policy effectiveness, and deployment scenarios spanning financial services, healthcare, and e-commerce. This work establishes decision governance as a foundational enterprise AI capability alongside model and API governance.

Keywords: AI Governance, Runtime Policy Enforcement, Risk Management, Autonomous Decision Systems, Anomaly Detection, Enterprise AI"""
    
    abstract_para = doc.add_paragraph(abstract_text)
    abstract_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    abstract_para.paragraph_format.space_before = Pt(6)
    abstract_para.paragraph_format.space_after = Pt(12)
    
    # Page break before TOC
    add_page_break(doc)
    
    # ===== TABLE OF CONTENTS =====
    toc_heading = doc.add_paragraph()
    toc_heading.style = 'Heading 1'
    toc_run = toc_heading.add_run("Table of Contents")
    toc_run.font.color.rgb = HEADING_COLOR
    
    toc_items = [
        "1. Introduction",
        "   1.1 The Autonomous AI Governance Gap",
        "   1.2 The Governance Paradox",
        "   1.3 Our Contribution: GlassBox",
        "   1.4 Evidence of Claims",
        "   1.5 Paper Structure",
        "2. Related Work",
        "3. Formal Model",
        "4. System Architecture",
        "5. Implementation",
        "6. Experimental Evaluation",
        "7. Enterprise Integration",
        "8. Limitations and Future Work",
        "9. Conclusion",
        "10. References",
        "Appendix A: Mermaid Diagrams",
        "Appendix B: Supplementary Materials",
    ]
    
    for item in toc_items:
        toc_para = doc.add_paragraph(item, style='List Number')
        toc_para.paragraph_format.left_indent = Inches(0.25 * (len(item) - len(item.lstrip())))
    
    add_page_break(doc)
    
    # ===== EXTRACT AND PARSE MAIN CONTENT =====
    # Split by section headers
    sections = re.split(r'\n## (\d+\.\d+.*?)$', content, flags=re.MULTILINE)
    
    # Process each section
    in_abstract = False
    for i, section in enumerate(sections):
        if i == 0:
            continue  # Skip header
        
        if i % 2 == 1:
            # Section heading
            heading = section.strip()
            if heading:
                heading_para = doc.add_paragraph()
                heading_run = heading_para.add_run(f"{heading}")
                heading_run.font.bold = True
                heading_run.font.size = Pt(13)
                heading_run.font.color.rgb = HEADING_COLOR
                heading_para.paragraph_format.space_before = Pt(12)
                heading_para.paragraph_format.space_after = Pt(6)
        else:
            # Section content
            lines = section.split('\n')
            for line in lines:
                line = line.rstrip()
                
                # Skip empty lines
                if not line:
                    continue
                
                # Skip Mermaid diagram code blocks (will handle separately)
                if line.strip().startswith('```mermaid'):
                    continue
                if line.strip().startswith('```'):
                    continue
                
                # Handle subsection headers
                if line.startswith('### '):
                    sub_heading = line.replace('### ', '').strip()
                    sub_para = doc.add_paragraph()
                    sub_run = sub_para.add_run(sub_heading)
                    sub_run.font.bold = True
                    sub_run.font.size = Pt(12)
                    sub_run.font.italic = True
                    sub_para.paragraph_format.space_before = Pt(8)
                    sub_para.paragraph_format.space_after = Pt(4)
                
                # Handle tables (simple markdown tables)
                elif line.strip().startswith('|'):
                    # Parse and add table
                    if '---' not in line:  # Skip separator rows
                        cols = [col.strip() for col in line.split('|')[1:-1]]
                        if cols and not line.startswith('| --- '):
                            # Add table row/header
                            pass
                
                # Handle code blocks
                elif line.startswith('    ') or line.startswith('\t'):
                    code_para = doc.add_paragraph(line.strip(), style='List Bullet')
                    code_para.paragraph_format.space_before = Pt(3)
                    code_para.paragraph_format.space_after = Pt(3)
                
                # Regular text
                else:
                    text_para = doc.add_paragraph(line)
                    text_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                    text_para.paragraph_format.space_after = Pt(6)
                    
                    # Ensure Times New Roman 12pt
                    for run in text_para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
    
    # ===== REFERENCES SECTION =====
    add_page_break(doc)
    ref_heading = doc.add_paragraph()
    ref_run = ref_heading.add_run("10. References")
    ref_run.font.bold = True
    ref_run.font.size = Pt(13)
    ref_run.font.color.rgb = HEADING_COLOR
    
    references = [
        "[1] Bauer, A., Leucker, M., & Schallhart, C. (2011). Runtime verification for LTL and MTL. ACM Transactions on Software Engineering and Methodology, 20(4), 1-64.",
        "[2] Bratman, M. E. (1987). Intention, Plans, and Practical Reason. Harvard University Press.",
        "[3] Chandola, V., Banerjee, A., & Kumar, V. (2009). Anomaly detection: A survey. ACM Computing Surveys, 41(3), 1-58.",
        "[4] EU (2018). General Data Protection Regulation (GDPR). Official Journal of the European Union.",
        "[5] EU (2024). EU AI Act. Official Journal of the European Union.",
        "[6] Gama, J., Sebastião, R., & Rodrigues, P. P. (2014). On evaluating stream learning algorithms. Machine Learning, 90(3), 317-346.",
        "[7] Google (2023). Responsible AI Toolkit. Google Cloud Documentation.",
        "[8] HHS (2013). HIPAA Security Rule. Department of Health and Human Services.",
        "[9] ISO (2022). Information Security Management Systems (ISO 27001:2022). International Organization for Standardization.",
        "[10] NIST (2018). Cybersecurity Framework (CSF). National Institute of Standards and Technology.",
        "[11] NIST (2023). AI Risk Management Framework (AI RMF). National Institute of Standards and Technology.",
        "[12] Netflix (2011). Hystrix: Latency and fault tolerance library. Netflix Open Source.",
        "[13] OpenAI (2022). Strategies for Building Robust AI Systems. OpenAI Blog.",
        "[14] RFC 6749 (2012). The OAuth 2.0 Authorization Framework. Internet Engineering Task Force.",
        "[15] Rossi, F., Venable, K. B., & Walsh, T. (2011). A Short Introduction to Preferences: Between Artificial Intelligence and Social Choice. Synthesis Lectures on Artificial Intelligence and Machine Learning.",
        "[16] Styra (2024). Rego: Policy Language. Styra Documentation.",
        "[17] Welford, B. P. (1962). Note on a method for calculating corrected sums of squares and products. Technometrics, 4(3), 419-420.",
        "[18] Wooldridge, M. J., & Jennings, N. R. (1995). Intelligent agents: Theory and practice. The Knowledge Engineering Review, 10(02), 115-152.",
    ]
    
    for ref in references:
        ref_para = doc.add_paragraph(ref)
        ref_para.paragraph_format.left_indent = Inches(0.5)
        ref_para.paragraph_format.first_line_indent = Inches(-0.5)
        ref_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        ref_para.paragraph_format.space_after = Pt(6)
    
    # Save document
    doc.save(output_file)
    print(f"✓ Word document created: {output_file}")
    return output_file

if __name__ == '__main__':
    markdown_file = r'c:\Akbar\Personal\AI Research Work\glassbox-agentic-governance\docs\GlassBox_Research_Paper_v2_0_Publication_Ready.md'
    output_file = r'c:\Akbar\Personal\AI Research Work\glassbox-agentic-governance\docs\Word\GlassBox_Publication_Ready_v2_0.docx'
    
    create_publication_word_doc(markdown_file, output_file)
    print(f"✓ Publication-ready Word document saved to: {output_file}")
